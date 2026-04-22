"""
ProposalCraft — Motor de Generación de Propuestas Comerciales
Motor genérico: sin hardcoding de marca. Lee toda la identidad desde BrandConfig.

Uso:
    from src.brand_loader import load_brand
    from src.proposal_engine import generate_proposal, ProposalData

    brand = load_brand()
    proposal = ProposalData(cliente="Empresa XYZ", ...)
    generate_proposal(brand, proposal)
"""

import os
from dataclasses import dataclass, field
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.brand_loader import BrandConfig, load_brand


# ══════════════════════════════════════════════════════════════════════════════
# DATOS DE LA PROPUESTA
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class ProposalData:
    """Contenido de la propuesta — independiente de la marca."""
    cliente: str
    titulo: str
    tagline: str
    resumen: str

    # Lista de tuplas: (label_modulo, duracion_tema, contenido_puntos)
    modulos: list = field(default_factory=list)

    # Lista de tuplas: (etiqueta, valor)
    detalles: list = field(default_factory=list)

    valor_total: str = "$ X.XXX.XXX"
    forma_pago: str  = "50% al confirmar · 50% al entregar"
    incluye: str     = "· Ítem 1\n· Ítem 2\n· Ítem 3"

    # Si se dejan en None, se calculan automáticamente desde hoy
    fecha_emision: str | None  = None
    fecha_vigencia: str | None = None

    output_filename: str | None = None  # Si None → "Propuesta_{cliente}.docx"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE DOCUMENTO (todos parametrizados con brand)
# ══════════════════════════════════════════════════════════════════════════════

def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(color))
    tcPr.append(shd)


def _set_cell_padding(cell, top=80, bottom=80, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_width(table, width_twips):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _set_col_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    if line:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)


def _add_run(para, text, brand: BrandConfig, bold=False, italic=False,
             size=9.5, color: RGBColor = None):
    if color is None:
        color = brand.color_body_text
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = brand.font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def _new_para(doc, brand: BrandConfig, text='', bold=False, italic=False,
              size=9.5, color: RGBColor = None,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              before=0, after=60, line=None):
    if color is None:
        color = brand.color_body_text
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, before=before, after=after, line=line)
    if text:
        _add_run(p, text, brand, bold=bold, italic=italic, size=size, color=color)
    return p


def _add_divider(doc, color_hex: str, sz=8, before=80, after=80):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _section_title(doc, brand: BrandConfig, text: str):
    """Título de sección con acento 1 de la marca."""
    _new_para(doc, brand, text, bold=True, size=8.5,
              color=brand.color_accent_1, before=0, after=28)


# ══════════════════════════════════════════════════════════════════════════════
# VARIABLES DE TEMA — layout estructural según preset activo
# ══════════════════════════════════════════════════════════════════════════════

def _theme_vars(brand: BrandConfig) -> dict:
    """
    Retorna las variables estructurales de color según brand.layout_preset.

    Preset "dark"  → OscuroPrestige: bloques oscuros dominantes, acento neon.
    Preset "light" → ClaroFormal:    todo claro, sin fondos oscuros, sin neon.
    Preset "dual"  → DualDinamico:   inversión invertida, col1 en accent_2.
    """
    preset = brand.layout_preset

    if preset == "light":
        return {
            "divider_top":         brand.hex_primary_dark,
            "divider_mid":         brand.hex_primary_dark,
            "divider_bot":         brand.hex_primary_dark,
            "mod_col1_bg":         brand.color_light_bg,
            "mod_col1_text":       brand.color_primary_dark,
            "mod_row_alt":         False,
            "inv_left_bg":         brand.color_light_bg,
            "inv_right_bg":        brand.color_white,
            "inv_value_color":     brand.color_primary_dark,
            "inv_right_label_color": brand.color_primary_dark,
            "inv_right_body_color":  brand.color_body_text,
        }

    if preset == "dual":
        return {
            "divider_top":         brand.hex_accent_2,
            "divider_mid":         brand.hex_accent_2,
            "divider_bot":         brand.hex_accent_2,
            "mod_col1_bg":         brand.color_accent_2,
            "mod_col1_text":       brand.color_white,
            "mod_row_alt":         True,
            "inv_left_bg":         brand.color_light_bg,
            "inv_right_bg":        brand.color_primary_dark,
            "inv_value_color":     brand.color_accent_1,
            "inv_right_label_color": brand.color_white,
            "inv_right_body_color":  RGBColor(0xCC, 0xCC, 0xCC),
        }

    # default → "dark" (OscuroPrestige)
    return {
        "divider_top":         brand.hex_accent_1,
        "divider_mid":         brand.hex_primary_dark,
        "divider_bot":         brand.hex_accent_1,
        "mod_col1_bg":         brand.color_primary_dark,
        "mod_col1_text":       brand.color_accent_1,
        "mod_row_alt":         True,
        "inv_left_bg":         brand.color_primary_dark,
        "inv_right_bg":        brand.color_light_bg,
        "inv_value_color":     brand.color_accent_1,
        "inv_right_label_color": brand.color_primary_dark,
        "inv_right_body_color":  brand.color_body_text,
    }


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def generate_proposal(brand: BrandConfig, proposal: ProposalData,
                      output_path: str | None = None) -> str:
    """
    Genera un documento Word (.docx) con la propuesta comercial.

    Args:
        brand:       Configuración de marca (cargada con load_brand())
        proposal:    Datos del contenido de la propuesta
        output_path: Ruta de salida. Si None, usa brand.output_dir/Propuesta_{cliente}.docx

    Returns:
        Ruta absoluta del archivo generado.
    """
    # ── Variables de tema ─────────────────────────────────────────────────────
    tv = _theme_vars(brand)

    # ── Fechas ────────────────────────────────────────────────────────────────
    MESES = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]

    def fmt_date(d: date) -> str:
        return f"{d.day} de {MESES[d.month - 1]} de {d.year}"

    today = date.today()
    fecha_emision  = proposal.fecha_emision  or fmt_date(today)
    fecha_vigencia = proposal.fecha_vigencia or fmt_date(today + timedelta(days=brand.validity_days))

    # ── Ruta de salida ────────────────────────────────────────────────────────
    if output_path is None:
        fname = proposal.output_filename or f"Propuesta_{proposal.cliente}.docx"
        output_path = os.path.join(brand.output_dir, fname)

    # ── Documento ─────────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    FULL = 9360  # twips útiles (página carta con márgenes 2.54)

    # ── 1. LOGO ───────────────────────────────────────────────────────────────
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_logo, before=0, after=40)

    logo_path = brand.logo_dark_path or brand.logo_light_path
    if logo_path:
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(4))
    else:
        _add_run(p_logo, brand.company_name, brand,
                 bold=True, size=16, color=brand.color_accent_1)

    _add_divider(doc, tv["divider_top"], sz=12, before=20, after=60)

    # ── 2. ENCABEZADO ─────────────────────────────────────────────────────────
    _new_para(doc, brand, 'PROPUESTA COMERCIAL', bold=True, size=8,
              color=brand.color_accent_1, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=0, after=6)
    _new_para(doc, brand, proposal.titulo, bold=True, size=17,
              color=brand.color_primary_dark, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=0, after=6)
    _new_para(doc, brand, proposal.tagline, italic=True, size=9.5,
              color=brand.color_mid_gray, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=0, after=10)
    _new_para(doc, brand, f'Dirigida a: {proposal.cliente}', bold=True, size=9,
              color=brand.color_body_text, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=0, after=80)

    _add_divider(doc, tv["divider_mid"], sz=6, before=0, after=70)

    # ── 3. RESUMEN EJECUTIVO ──────────────────────────────────────────────────
    _section_title(doc, brand, 'RESUMEN EJECUTIVO')
    _new_para(doc, brand, proposal.resumen, size=9.5,
              color=brand.color_body_text, before=0, after=80, line=264)

    _add_divider(doc, tv["divider_mid"], sz=6, before=0, after=70)

    # ── 4. ALCANCE / MÓDULOS ──────────────────────────────────────────────────
    if proposal.modulos:
        _section_title(doc, brand, 'ALCANCE DEL SERVICIO')

        W0, W1, W2 = 1100, 2200, 6060
        tmod = doc.add_table(rows=1 + len(proposal.modulos), cols=3)
        tmod.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tmod, FULL)
        _no_borders(tmod)

        # Header de tabla
        hr = tmod.rows[0]
        for ci, (txt, w) in enumerate(zip(['', 'DURACIÓN / TEMA', 'CONTENIDO'], [W0, W1, W2])):
            cell = hr.cells[ci]
            _set_cell_bg(cell, tv["mod_col1_bg"])
            _set_col_width(cell, w)
            _set_cell_padding(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _para_spacing(p, before=0, after=0)
            if txt:
                _add_run(p, txt, brand, bold=True, size=8, color=brand.color_white)

        # Filas de módulos
        for ridx, (mod, dur, cont) in enumerate(proposal.modulos):
            row = tmod.rows[ridx + 1]
            bg = brand.color_light_bg if (tv["mod_row_alt"] and ridx % 2 == 0) else brand.color_white

            c0 = row.cells[0]
            _set_cell_bg(c0, tv["mod_col1_bg"])
            _set_col_width(c0, W0)
            _set_cell_padding(c0)
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _para_spacing(p0, before=0, after=0)
            _add_run(p0, mod, brand, bold=True, size=7.5, color=tv["mod_col1_text"])

            c1 = row.cells[1]
            _set_cell_bg(c1, bg)
            _set_col_width(c1, W1)
            _set_cell_padding(c1)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _para_spacing(p1, before=0, after=0)
            parts = dur.split('\n', 1)
            _add_run(p1, parts[0], brand, bold=True, size=9, color=brand.color_accent_1)
            if len(parts) > 1:
                _add_run(p1, '\n' + parts[1], brand, bold=False, size=8,
                         color=brand.color_body_text)

            c2 = row.cells[2]
            _set_cell_bg(c2, bg)
            _set_col_width(c2, W2)
            _set_cell_padding(c2)
            c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _para_spacing(p2, before=0, after=0)
            _add_run(p2, cont, brand, size=8.5, color=brand.color_body_text)

    # ── 5. DETALLES ───────────────────────────────────────────────────────────
    if proposal.detalles:
        _add_divider(doc, tv["divider_mid"], sz=6, before=80, after=70)
        _section_title(doc, brand, 'DETALLES DEL PROGRAMA')

        WD0, WD1 = 2500, 6860
        tdet = doc.add_table(rows=len(proposal.detalles), cols=2)
        tdet.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_width(tdet, FULL)
        _no_borders(tdet)

        for ridx, (lbl, val) in enumerate(proposal.detalles):
            row = tdet.rows[ridx]
            bg = brand.color_light_bg if ridx % 2 == 0 else brand.color_white
            c0 = row.cells[0]
            _set_cell_bg(c0, bg)
            _set_col_width(c0, WD0)
            _set_cell_padding(c0, top=60, bottom=60, left=100, right=80)
            _para_spacing(c0.paragraphs[0], before=0, after=0)
            _add_run(c0.paragraphs[0], lbl, brand, bold=True, size=8.5,
                     color=brand.color_primary_dark)
            c1 = row.cells[1]
            _set_cell_bg(c1, bg)
            _set_col_width(c1, WD1)
            _set_cell_padding(c1, top=60, bottom=60, left=100, right=80)
            _para_spacing(c1.paragraphs[0], before=0, after=0)
            _add_run(c1.paragraphs[0], val, brand, size=8.5,
                     color=brand.color_body_text)

    # ── 6. INVERSIÓN ──────────────────────────────────────────────────────────
    _add_divider(doc, tv["divider_mid"], sz=6, before=80, after=70)
    _section_title(doc, brand, 'INVERSIÓN')

    WI0, WI1 = 3800, 5560
    tinv = doc.add_table(rows=1, cols=2)
    tinv.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(tinv, FULL)
    _no_borders(tinv)

    ci_cell = tinv.rows[0].cells[0]
    _set_cell_bg(ci_cell, tv["inv_left_bg"])
    _set_col_width(ci_cell, WI0)
    _set_cell_padding(ci_cell, top=120, bottom=120, left=120, right=120)
    ci_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = ci_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=0, after=0)
    _add_run(p, 'INVERSIÓN TOTAL\n', brand, bold=True, size=8,
             color=brand.color_mid_gray)
    _add_run(p, proposal.valor_total, brand, bold=True, size=17,
             color=tv["inv_value_color"])

    cd = tinv.rows[0].cells[1]
    _set_cell_bg(cd, tv["inv_right_bg"])
    _set_col_width(cd, WI1)
    _set_cell_padding(cd, top=100, bottom=100, left=120, right=100)
    cd.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cd.paragraphs[0]
    _para_spacing(p, before=0, after=20)
    _add_run(p, 'Forma de pago\n', brand, bold=True, size=8.5,
             color=tv["inv_right_label_color"])
    _add_run(p, proposal.forma_pago, brand, size=8.5, color=tv["inv_right_body_color"])
    p2 = cd.add_paragraph()
    _para_spacing(p2, before=40, after=0)
    _add_run(p2, 'Incluye\n', brand, bold=True, size=8.5,
             color=tv["inv_right_label_color"])
    _add_run(p2, proposal.incluye, brand, size=8.5, color=tv["inv_right_body_color"])

    # ── 7. VIGENCIA + FIRMA ───────────────────────────────────────────────────
    _add_divider(doc, tv["divider_mid"], sz=6, before=80, after=70)

    WV0, WV1 = 4680, 4680
    tvf = doc.add_table(rows=1, cols=2)
    tvf.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(tvf, FULL)
    _no_borders(tvf)

    cv = tvf.rows[0].cells[0]
    _set_cell_bg(cv, brand.color_light_bg)
    _set_col_width(cv, WV0)
    _set_cell_padding(cv)
    p = cv.paragraphs[0]
    _para_spacing(p, before=0, after=16)
    _add_run(p, 'VIGENCIA DE LA PROPUESTA\n', brand, bold=True, size=8.5,
             color=brand.color_primary_dark)
    _add_run(p, f'Emisión:         {fecha_emision}\n', brand, size=8.5,
             color=brand.color_body_text)
    _add_run(p, f'Válida hasta:  {fecha_vigencia}', brand, size=8.5,
             color=brand.color_body_text)

    cf = tvf.rows[0].cells[1]
    _set_cell_bg(cf, brand.color_white)
    _set_col_width(cf, WV1)
    _set_cell_padding(cf)
    p = cf.paragraphs[0]
    _para_spacing(p, before=0, after=16)
    _add_run(p, 'PROPONENTE\n', brand, bold=True, size=8.5,
             color=brand.color_primary_dark)
    _add_run(p, f'{brand.proponent_name}\n', brand, bold=True, size=9,
             color=brand.color_body_text)
    _add_run(p, f'{brand.proponent_id_full}\n\n', brand, size=8.5,
             color=brand.color_body_text)
    _add_run(p, '___________________________\nFirma', brand, size=8,
             color=brand.color_mid_gray)

    # ── 8. PIE DE PÁGINA ─────────────────────────────────────────────────────
    _add_divider(doc, tv["divider_bot"], sz=6, before=60, after=20)
    footer_text = brand.proponent_full
    if brand.proponent_id_full:
        footer_text += f' · {brand.proponent_id_full}'
    _new_para(doc, brand, footer_text, size=7.5, color=brand.color_mid_gray,
              align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    # ── 9. BANNER DE CIERRE (opcional) ───────────────────────────────────────
    if brand.banner_path:
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p_banner, before=40, after=0)
        run_b = p_banner.add_run()
        run_b.add_picture(brand.banner_path, width=Cm(17))

    # ── Guardar ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Propuesta generada: {output_path}")
    return output_path
